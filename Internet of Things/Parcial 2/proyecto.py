#! /usr/bin/env python3
# -*- coding: utf-8 -*-

import csv
from datetime import date, datetime
import os
from pathlib import Path
import customtkinter as ctk
from tkinter import messagebox, filedialog
import base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Imports para OAuth2
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

# === CONFIGURACIÓN ===
ARCHIVO_INFO = "res/info.csv"
ARCHIVO_PLANTILLA = "res/oficio.txt"
CARPETA_SALIDA = "constancias"

# Configuración OAuth2 para Gmail
SCOPES = ['https://www.googleapis.com/auth/gmail.send']
TOKEN_PATH = 'res/token.json'
CREDENTIALS_PATH = 'res/credentials.json'

# Configuración de la apariencia
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


# === FUNCIONES AUXILIARES ===

def calcular_edad(fecha_nacimiento):
    """Calcula la edad a partir de una fecha de nacimiento (YYYY-MM-DD)."""
    try:
        nacimiento = datetime.strptime(fecha_nacimiento, "%Y-%m-%d").date()
        hoy = date.today()
        edad = hoy.year - nacimiento.year - ((hoy.month, hoy.day) < (nacimiento.month, nacimiento.day))
        return edad
    except:
        return 0


def fecha_actual_formateada():
    """Devuelve la fecha actual en formato: '21 de octubre de 2025'."""
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    hoy = date.today()
    return f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"


def cargar_personas():
    """Carga las personas desde el archivo CSV."""
    personas = []
    if os.path.exists(ARCHIVO_INFO):
        with open(ARCHIVO_INFO, "r", encoding="utf-8") as f:
            lector = csv.DictReader(f)
            for persona in lector:
                personas.append(persona)
    return personas


def guardar_personas(personas):
    """Guarda las personas en el archivo CSV."""
    if not personas:
        return
    
    os.makedirs("res", exist_ok=True)
    fieldnames = personas[0].keys()
    
    with open(ARCHIVO_INFO, "w", encoding="utf-8", newline='') as f:
        escritor = csv.DictWriter(f, fieldnames=fieldnames)
        escritor.writeheader()
        escritor.writerows(personas)


def generar_oficio_texto(persona, plantilla):
    """Genera el texto del oficio para una persona."""
    persona_copy = persona.copy()
    persona_copy["edad"] = calcular_edad(persona["fechaNacimiento"])
    persona_copy["fechaActual"] = fecha_actual_formateada()
    
    numero_int = persona["numeroInt"].strip()
    persona_copy["numeroIntTexto"] = f" Int. {numero_int}" if numero_int else ""
    
    return plantilla.format(**persona_copy)


def generar_txt(personas_seleccionadas):
    """Genera archivos de texto plano."""
    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    
    with open(ARCHIVO_PLANTILLA, "r", encoding="utf-8") as f:
        plantilla = f.read()
    
    archivos_generados = []
    
    for persona in personas_seleccionadas:
        oficio_texto = generar_oficio_texto(persona, plantilla)
        nombre_archivo = f"oficio_{persona['registro']}.txt"
        ruta_salida = os.path.join(CARPETA_SALIDA, nombre_archivo)
        
        with open(ruta_salida, "w", encoding="utf-8") as salida:
            salida.write(oficio_texto)
        
        archivos_generados.append(ruta_salida)
    
    return archivos_generados


def generar_docx(personas_seleccionadas):
    """Genera archivos Word (.docx) con logo, marco y fondo."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        messagebox.showerror("Error", "Necesitas instalar python-docx: pip install python-docx")
        return []
    
    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    
    with open(ARCHIVO_PLANTILLA, "r", encoding="utf-8") as f:
        plantilla = f.read()
    
    archivos_generados = []
    logo_path = "res/logo.png"
    fondo_path = "res/fondo.png"
    
    for persona in personas_seleccionadas:
        oficio_texto = generar_oficio_texto(persona, plantilla)
        
        doc = Document()
        
        # Configurar márgenes
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Agregar imagen de fondo si existe
        if os.path.exists(fondo_path):
            def add_watermark(section, image_path):
                """Agrega una imagen de fondo/marca de agua a la sección."""
                try:
                    page_width = section.page_width
                    page_height = section.page_height
                    
                    header = section.header
                    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    
                    run = header_para.add_run()
                    picture = run.add_picture(image_path)
                    
                    picture.width = page_width
                    picture.height = page_height
                    
                    inline = run._element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                    if inline is not None:
                        anchor = OxmlElement('wp:anchor')
                        anchor.set('behindDoc', '1')
                        anchor.set('locked', '0')
                        anchor.set('layoutInCell', '1')
                        anchor.set('allowOverlap', '1')
                        
                        for attr in inline.attrib:
                            if attr not in ['distT', 'distB', 'distL', 'distR']:
                                anchor.set(attr, inline.get(attr))
                        
                        for child in inline:
                            anchor.append(child)
                        
                        parent = inline.getparent()
                        parent.replace(inline, anchor)
                        
                        positionH = OxmlElement('wp:positionH')
                        positionH.set('relativeFrom', 'page')
                        positionH_offset = OxmlElement('wp:posOffset')
                        positionH_offset.text = '0'
                        positionH.append(positionH_offset)
                        anchor.append(positionH)
                        
                        positionV = OxmlElement('wp:positionV')
                        positionV.set('relativeFrom', 'page')
                        positionV_offset = OxmlElement('wp:posOffset')
                        positionV_offset.text = '0'
                        positionV.append(positionV_offset)
                        anchor.append(positionV)
                except Exception as e:
                    print(f"Error al agregar fondo: {e}")
            
            add_watermark(section, fondo_path)
        
        # Agregar logo si existe
        if os.path.exists(logo_path):
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            doc.add_paragraph()
        
        # Agregar el contenido del oficio
        for linea in oficio_texto.split('\n'):
            if linea.strip():
                p = doc.add_paragraph(linea.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
            else:
                doc.add_paragraph()
        
        # Agregar marco/borde
        def add_page_border(section):
            """Agrega un borde a la página."""
            sectPr = section._sectPr
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '24')
                border.set(qn('w:space'), '24')
                border.set(qn('w:color'), '000000')
                pgBorders.append(border)
            
            sectPr.append(pgBorders)
        
        for section in doc.sections:
            add_page_border(section)
        
        nombre_archivo = f"oficio_{persona['registro']}.docx"
        ruta_salida = os.path.join(CARPETA_SALIDA, nombre_archivo)
        doc.save(ruta_salida)
        
        archivos_generados.append(ruta_salida)
    
    return archivos_generados


def generar_pdf(personas_seleccionadas):
    """Genera archivos PDF con formato profesional, texto justificado y fondo."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_JUSTIFY
        from reportlab.lib import colors
    except ImportError:
        messagebox.showerror("Error", "Necesitas instalar reportlab: pip install reportlab")
        return []
    
    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    
    with open(ARCHIVO_PLANTILLA, "r", encoding="utf-8") as f:
        plantilla = f.read()
    
    archivos_generados = []
    logo_path = "res/logo.png"
    sello_path = "res/sello.jpg"
    fondo_path = "res/fondo.png"
    
    for persona in personas_seleccionadas:
        oficio_texto = generar_oficio_texto(persona, plantilla)
        
        nombre_archivo = f"oficio_{persona['registro']}.pdf"
        ruta_salida = os.path.join(CARPETA_SALIDA, nombre_archivo)
        
        def add_page_decorations(canvas, doc):
            """Agrega fondo y marco/borde a cada página."""
            canvas.saveState()
            
            if os.path.exists(fondo_path):
                try:
                    canvas.drawImage(
                        fondo_path,
                        0, 0,
                        width=letter[0],
                        height=letter[1],
                        preserveAspectRatio=False,
                        mask='auto'
                    )
                except Exception as e:
                    print(f"Error al cargar fondo: {e}")
            
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(2)
            margin = 0.5 * inch
            canvas.rect(
                margin, 
                margin, 
                letter[0] - 2*margin, 
                letter[1] - 2*margin
            )
            canvas.restoreState()
        
        doc = SimpleDocTemplate(
            ruta_salida,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )
        
        styles = getSampleStyleSheet()
        
        estilo_justificado = ParagraphStyle(
            'CustomJustified',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            alignment=TA_JUSTIFY,
            fontName='Helvetica',
            spaceAfter=0,
            spaceBefore=0
        )
        
        contenido = []
        
        if os.path.exists(logo_path):
            img = Image(logo_path, width=1.5*inch, height=1.5*inch, kind='proportional')
            img.hAlign = 'CENTER'
            contenido.append(img)
            contenido.append(Spacer(1, 0.3*inch))
        
        lineas = oficio_texto.split('\n')
        
        indice_firma = -1
        for i, linea in enumerate(lineas):
            if '_____' in linea or '____' in linea:
                indice_firma = i
                break
        
        if indice_firma >= 0:
            for i in range(indice_firma):
                linea = lineas[i]
                if linea.strip():
                    contenido.append(Paragraph(linea.strip(), estilo_justificado))
                else:
                    contenido.append(Spacer(1, 0.15*inch))
            
            datos_tabla = []
            
            texto_firma = []
            for i in range(indice_firma, len(lineas)):
                if lineas[i].strip():
                    texto_firma.append(Paragraph(lineas[i].strip(), estilo_justificado))
                else:
                    texto_firma.append(Spacer(1, 0.1*inch))
            
            if os.path.exists(sello_path):
                sello_img = Image(sello_path, width=2*inch, height=2*inch, kind='proportional')
            else:
                sello_img = Paragraph("", estilo_justificado)
            
            datos_tabla = [[texto_firma, sello_img]]
            
            tabla = Table(datos_tabla, colWidths=[3.5*inch, 2.5*inch])
            tabla.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (1, 0), (1, 0), 'CENTER'),
            ]))
            
            contenido.append(tabla)
        else:
            for linea in lineas:
                if linea.strip():
                    contenido.append(Paragraph(linea.strip(), estilo_justificado))
                else:
                    contenido.append(Spacer(1, 0.15*inch))
        
        doc.build(contenido, onFirstPage=add_page_decorations, onLaterPages=add_page_decorations)
        archivos_generados.append(ruta_salida)
    
    return archivos_generados


# === FUNCIONES OAUTH2 PARA GMAIL ===

def get_gmail_service():
    """Obtiene el servicio de Gmail usando OAuth2."""
    creds = None
    
    # Verificar si existe el token guardado
    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)
    
    # Si no hay credenciales válidas, pedir autorización
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(CREDENTIALS_PATH):
                raise FileNotFoundError(
                    f"No se encontró el archivo {CREDENTIALS_PATH}\n\n"
                    "Debes crear credenciales OAuth2 en Google Cloud Console:\n"
                    "1. Ve a https://console.cloud.google.com/\n"
                    "2. Habilita Gmail API\n"
                    "3. Crea credenciales OAuth 2.0\n"
                    "4. Descarga el JSON y guárdalo como res/credentials.json"
                )
            
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_PATH, SCOPES)
            creds = flow.run_local_server(port=0)
        
        # Guardar las credenciales para la próxima vez
        with open(TOKEN_PATH, 'w') as token:
            token.write(creds.to_json())
    
    return build('gmail', 'v1', credentials=creds)


def enviar_email_oauth2(destinatario, asunto, cuerpo, archivos_adjuntos):
    """Envía un email usando OAuth2 de Gmail."""
    try:
        service = get_gmail_service()
        
        # Crear el mensaje
        message = MIMEMultipart()
        message['To'] = destinatario
        message['Subject'] = asunto
        
        # Agregar el cuerpo del mensaje
        message.attach(MIMEText(cuerpo, 'plain', 'utf-8'))
        
        # Adjuntar archivos
        for archivo_path in archivos_adjuntos:
            with open(archivo_path, 'rb') as f:
                parte = MIMEBase('application', 'octet-stream')
                parte.set_payload(f.read())
                encoders.encode_base64(parte)
                parte.add_header(
                    'Content-Disposition',
                    f'attachment; filename={os.path.basename(archivo_path)}'
                )
                message.attach(parte)
        
        # Codificar el mensaje
        raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')
        
        # Enviar el mensaje
        send_message = service.users().messages().send(
            userId='me',
            body={'raw': raw_message}
        ).execute()
        
        return True, f"Mensaje enviado con ID: {send_message['id']}"
        
    except HttpError as error:
        return False, f"Error de Gmail API: {error}"
    except FileNotFoundError as error:
        return False, str(error)
    except Exception as error:
        return False, f"Error: {str(error)}"


# === INTERFAZ GRÁFICA ===

class AplicacionOficios(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        self.title("Sistema de Gestión de Oficios")
        self.geometry("1000x700")
        
        self.personas = cargar_personas()
        
        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)
        
        self.tab_lista = self.tabview.add("Lista de Personas")
        self.tab_alta = self.tabview.add("Alta")
        self.tab_modificar = self.tabview.add("Modificar")
        self.tab_generar = self.tabview.add("Generar Oficios")
        
        self.crear_tab_lista()
        self.crear_tab_alta()
        self.crear_tab_modificar()
        self.crear_tab_generar()
    
    def crear_tab_lista(self):
        """Crea la interfaz de lista de personas."""
        # Frame superior con botones
        frame_botones = ctk.CTkFrame(self.tab_lista)
        frame_botones.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkButton(frame_botones, text="🔄 Recargar", command=self.recargar_lista).pack(side="left", padx=5)
        ctk.CTkButton(frame_botones, text="🗑️ Eliminar Seleccionado", command=self.eliminar_persona).pack(side="left", padx=5)
        
        # Frame con scrollbar
        frame_scroll = ctk.CTkScrollableFrame(self.tab_lista)
        frame_scroll.pack(fill="both", expand=True, padx=10, pady=5)
        
        self.frame_lista = frame_scroll
        self.actualizar_lista()

    def actualizar_lista(self):
        """Actualiza la lista de personas."""
        for widget in self.frame_lista.winfo_children():
            widget.destroy()
        
        if not self.personas:
            ctk.CTkLabel(self.frame_lista, text="No hay personas registradas", 
                        font=("Arial", 14)).pack(pady=20)
            return
        
        # Definir anchos fijos para cada columna (en píxeles)
        anchos_columnas = [100, 250, 120, 200, 150, 100]
        
        # Encabezados
        headers = ["Registro", "Nombre", "Fecha Nac.", "Calle", "Colonia", "CP"]
        frame_header = ctk.CTkFrame(self.frame_lista)
        frame_header.pack(fill="x", pady=(0, 5))
        
        for i, header in enumerate(headers):
            label = ctk.CTkLabel(
                frame_header, 
                text=header, 
                font=("Arial", 12, "bold"),
                width=anchos_columnas[i],
                anchor="w"  # Alinear a la izquierda
            )
            label.grid(row=0, column=i, padx=5, pady=5, sticky="w")
    
        # Datos
        for idx, persona in enumerate(self.personas):
            frame_persona = ctk.CTkFrame(self.frame_lista)
            frame_persona.pack(fill="x", pady=2)
            
            datos = [
                persona.get('registro', ''),
                persona.get('nombre', ''),
                persona.get('fechaNacimiento', ''),
                persona.get('calle', ''),
                persona.get('colonia', ''),
                persona.get('codigoPostal', '')
            ]
            
            for i, dato in enumerate(datos):
                label = ctk.CTkLabel(
                    frame_persona, 
                    text=dato,
                    width=anchos_columnas[i],
                    anchor="w"  # Alinear a la izquierda
                )
                label.grid(row=0, column=i, padx=5, pady=5, sticky="w")
    
    def recargar_lista(self):
        """Recarga la lista desde el archivo."""
        self.personas = cargar_personas()
        self.actualizar_lista()
        messagebox.showinfo("Éxito", "Lista recargada correctamente")
    
    def eliminar_persona(self):
        """Elimina una persona."""
        if not self.personas:
            messagebox.showwarning("Aviso", "No hay personas para eliminar")
            return
        
        ventana = ctk.CTkToplevel(self)
        ventana.title("Eliminar Persona")
        ventana.geometry("400x200")
        
        ctk.CTkLabel(ventana, text="Selecciona el registro a eliminar:").pack(pady=10)
        
        registros = [f"{p['registro']} - {p['nombre']}" for p in self.personas]
        combo = ctk.CTkComboBox(ventana, values=registros, width=350)
        combo.pack(pady=10)
        combo.set(registros[0] if registros else "")
        
        def confirmar_eliminacion():
            seleccion = combo.get()
            if not seleccion:
                return
            
            registro = seleccion.split(" - ")[0]
            self.personas = [p for p in self.personas if p['registro'] != registro]
            guardar_personas(self.personas)
            self.actualizar_lista()
            ventana.destroy()
            messagebox.showinfo("Éxito", f"Persona con registro {registro} eliminada")
        
        ctk.CTkButton(ventana, text="Eliminar", command=confirmar_eliminacion).pack(pady=10)
    
    def crear_tab_alta(self):
        """Crea la interfaz de alta de personas."""
        frame_form = ctk.CTkScrollableFrame(self.tab_alta)
        frame_form.pack(fill="both", expand=True, padx=20, pady=20)
        
        campos = [
            ("Registro", "registro"),
            ("Nombre(s)", "nombre"),
            ("Apellido Paterno", "apellido1"),
            ("Apellido Materno", "apellido2"),
            ("Carrera", "carrera"),
            ("Semestre", "semestre"),
            ("Fecha de Nacimiento (YYYY-MM-DD)", "fechaNacimiento"),
            ("Calle", "calle"),
            ("Número Ext.", "numeroExt"),
            ("Número Int.", "numeroInt"),
            ("Colonia", "colonia"),
            ("Código Postal", "codigoPostal"),
            ("Municipio", "municipio"),
            ("Estado", "estado"),
            ("Teléfono", "telefono"),
            ("Correo Electrónico", "correoElectronico")
        ]
        
        self.entries_alta = {}
        
        for i, (label, key) in enumerate(campos):
            ctk.CTkLabel(frame_form, text=label + ":", font=("Arial", 12)).grid(
                row=i, column=0, sticky="w", pady=5, padx=5)
            entry = ctk.CTkEntry(frame_form, width=300)
            entry.grid(row=i, column=1, pady=5, padx=5)
            self.entries_alta[key] = entry
        
        ctk.CTkButton(frame_form, text="💾 Guardar Persona", 
                     command=self.guardar_nueva_persona, height=40).grid(
            row=len(campos), column=0, columnspan=2, pady=20)
    
    def guardar_nueva_persona(self):
        """Guarda una nueva persona."""
        nueva_persona = {}
        
        for key, entry in self.entries_alta.items():
            valor = entry.get().strip()
            if not valor and key not in ["numeroInt"]:
                messagebox.showwarning("Aviso", f"El campo {key} es obligatorio")
                return
            nueva_persona[key] = valor
        
        if any(p['registro'] == nueva_persona['registro'] for p in self.personas):
            messagebox.showerror("Error", "Ya existe una persona con ese registro")
            return
        
        self.personas.append(nueva_persona)
        guardar_personas(self.personas)
        
        for entry in self.entries_alta.values():
            entry.delete(0, 'end')
        
        self.actualizar_lista()
        messagebox.showinfo("Éxito", "Persona agregada correctamente")
    
    def crear_tab_modificar(self):
        """Crea la interfaz de modificación."""
        frame_principal = ctk.CTkFrame(self.tab_modificar)
        frame_principal.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(frame_principal, text="Seleccionar Persona:", 
                    font=("Arial", 14, "bold")).pack(pady=10)
        
        self.combo_modificar = ctk.CTkComboBox(frame_principal, width=400,
                                               command=self.cargar_datos_modificar)
        self.combo_modificar.pack(pady=5)
        
        self.frame_campos_modificar = ctk.CTkScrollableFrame(frame_principal, height=400)
        self.frame_campos_modificar.pack(fill="both", expand=True, pady=10)
        
        ctk.CTkButton(frame_principal, text="💾 Actualizar Datos", 
                     command=self.actualizar_persona, height=40).pack(pady=10)
        
        self.actualizar_combo_modificar()
    
    def actualizar_combo_modificar(self):
        """Actualiza el combo de personas para modificar."""
        if self.personas:
            registros = [f"{p['registro']} - {p['nombre']}" for p in self.personas]
            self.combo_modificar.configure(values=registros)
            if registros:
                self.combo_modificar.set(registros[0])
                self.cargar_datos_modificar(registros[0])
    
    def cargar_datos_modificar(self, seleccion):
        """Carga los datos de la persona seleccionada."""
        for widget in self.frame_campos_modificar.winfo_children():
            widget.destroy()
        
        if not seleccion:
            return
        
        registro = seleccion.split(" - ")[0]
        persona = next((p for p in self.personas if p['registro'] == registro), None)
        
        if not persona:
            return
        
        self.entries_modificar = {}
        
        campos = [
            ("Registro", "registro"),
            ("Nombre(s)", "nombre"),
            ("Apellido Paterno", "apellido1"),
            ("Apellido Materno", "apellido2"),
            ("Carrera", "carrera"),
            ("Semestre", "semestre"),
            ("Fecha de Nacimiento (YYYY-MM-DD)", "fechaNacimiento"),
            ("Calle", "calle"),
            ("Número Ext.", "numeroExt"),
            ("Número Int.", "numeroInt"),
            ("Colonia", "colonia"),
            ("Código Postal", "codigoPostal"),
            ("Municipio", "municipio"),
            ("Estado", "estado"),
            ("Teléfono", "telefono"),
            ("Correo Electrónico", "correoElectronico")
        ]
        
        for i, (label, key) in enumerate(campos):
            ctk.CTkLabel(self.frame_campos_modificar, text=label + ":", 
                        font=("Arial", 12)).grid(row=i, column=0, sticky="w", pady=5, padx=5)
            entry = ctk.CTkEntry(self.frame_campos_modificar, width=300)
            entry.grid(row=i, column=1, pady=5, padx=5)
            entry.insert(0, persona.get(key, ''))
            
            if key == "registro":
                entry.configure(state="disabled")
            
            self.entries_modificar[key] = entry
    
    def actualizar_persona(self):
        """Actualiza los datos de una persona."""
        seleccion = self.combo_modificar.get()
        if not seleccion:
            return
        
        registro = seleccion.split(" - ")[0]
        
        for persona in self.personas:
            if persona['registro'] == registro:
                for key, entry in self.entries_modificar.items():
                    if key != "registro":
                        persona[key] = entry.get().strip()
                break
        
        guardar_personas(self.personas)
        self.actualizar_lista()
        self.actualizar_combo_modificar()
        messagebox.showinfo("Éxito", "Datos actualizados correctamente")
    
    def crear_tab_generar(self):
        """Crea la interfaz de generación de oficios."""
        frame_principal = ctk.CTkFrame(self.tab_generar)
        frame_principal.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(frame_principal, text="Seleccionar Personas:", 
                    font=("Arial", 14, "bold")).pack(pady=10)
        
        frame_seleccion = ctk.CTkFrame(frame_principal)
        frame_seleccion.pack(fill="x", pady=5)
        
        ctk.CTkButton(frame_seleccion, text="✅ Seleccionar Todos", 
                     command=self.seleccionar_todos).pack(side="left", padx=5)
        ctk.CTkButton(frame_seleccion, text="❌ Deseleccionar Todos", 
                     command=self.deseleccionar_todos).pack(side="left", padx=5)
        
        self.frame_checks = ctk.CTkScrollableFrame(frame_principal, height=300)
        self.frame_checks.pack(fill="both", expand=True, pady=10)
        
        self.checkboxes = []
        self.actualizar_checkboxes()
        
        frame_opciones = ctk.CTkFrame(frame_principal)
        frame_opciones.pack(fill="x", pady=10)
        
        ctk.CTkLabel(frame_opciones, text="Formato:", font=("Arial", 12)).pack(side="left", padx=5)
        
        self.formato_var = ctk.StringVar(value="TXT")
        
        ctk.CTkRadioButton(frame_opciones, text="TXT", variable=self.formato_var, 
                          value="TXT").pack(side="left", padx=5)
        ctk.CTkRadioButton(frame_opciones, text="DOCX", variable=self.formato_var, 
                          value="DOCX").pack(side="left", padx=5)
        ctk.CTkRadioButton(frame_opciones, text="PDF", variable=self.formato_var, 
                          value="PDF").pack(side="left", padx=5)
        
        frame_botones = ctk.CTkFrame(frame_principal)
        frame_botones.pack(fill="x", pady=10)
        
        ctk.CTkButton(frame_botones, text="Generar Oficios", 
                     command=self.generar_oficios, height=40).pack(side="left", padx=5, expand=True, fill="x")
        ctk.CTkButton(frame_botones, text="Enviar por Gmail", 
                     command=self.enviar_email_oauth2, height=40, 
                     fg_color="#DB4437", hover_color="#C23321").pack(side="left", padx=5, expand=True, fill="x")
    
    def actualizar_checkboxes(self):
        """Actualiza la lista de checkboxes."""
        for widget in self.frame_checks.winfo_children():
            widget.destroy()
        
        self.checkboxes = []
        
        if not self.personas:
            ctk.CTkLabel(self.frame_checks, text="No hay personas registradas").pack(pady=20)
            return
        
        for persona in self.personas:
            var = ctk.BooleanVar(value=False)
            texto = f"{persona['registro']} - {persona['nombre']}"
            check = ctk.CTkCheckBox(self.frame_checks, text=texto, variable=var)
            check.pack(anchor="w", pady=2)
            self.checkboxes.append((var, persona))
    
    def seleccionar_todos(self):
        """Selecciona todos los checkboxes."""
        for var, _ in self.checkboxes:
            var.set(True)
    
    def deseleccionar_todos(self):
        """Deselecciona todos los checkboxes."""
        for var, _ in self.checkboxes:
            var.set(False)
    
    def generar_oficios(self):
        """Genera los oficios según las opciones seleccionadas."""
        personas_seleccionadas = [persona for var, persona in self.checkboxes if var.get()]
        
        if not personas_seleccionadas:
            messagebox.showwarning("Aviso", "Selecciona al menos una persona")
            return
        
        if not os.path.exists(ARCHIVO_PLANTILLA):
            messagebox.showerror("Error", f"No se encuentra el archivo plantilla: {ARCHIVO_PLANTILLA}")
            return
        
        formato = self.formato_var.get()
        
        try:
            if formato == "TXT":
                archivos = generar_txt(personas_seleccionadas)
            elif formato == "DOCX":
                archivos = generar_docx(personas_seleccionadas)
            elif formato == "PDF":
                archivos = generar_pdf(personas_seleccionadas)
            
            messagebox.showinfo("Éxito", 
                              f"Se generaron {len(archivos)} oficios en formato {formato}\n"
                              f"Ubicación: {CARPETA_SALIDA}/")
        except Exception as e:
            messagebox.showerror("Error", f"Error al generar oficios: {str(e)}")
    
    def enviar_email_oauth2(self):
        """Envía oficios por email usando OAuth2 de Gmail."""
        personas_seleccionadas = [persona for var, persona in self.checkboxes if var.get()]
        
        if not personas_seleccionadas:
            messagebox.showwarning("Aviso", "Selecciona al menos una persona")
            return
        
        # Ventana de configuración
        ventana = ctk.CTkToplevel(self)
        ventana.title("Enviar por Gmail")
        ventana.geometry("550x400")
        
        ctk.CTkLabel(ventana, text="Enviar por Gmail", 
                    font=("Arial", 16, "bold")).pack(pady=15)
        
        # Instrucciones
        frame_info = ctk.CTkFrame(ventana)
        frame_info.pack(fill="x", padx=20, pady=10)
        
        info_text = (
            "Para enviar los oficios generados por Gmail, ingresa el destinatario y asunto.\n\n"
        )
        
        ctk.CTkLabel(frame_info, text=info_text, justify="left", 
                    wraplength=500).pack(padx=10, pady=10)
        
        # Frame para campos
        frame_campos = ctk.CTkFrame(ventana)
        frame_campos.pack(fill="both", expand=True, padx=20, pady=10)
        
        ctk.CTkLabel(frame_campos, text="Destinatario:").grid(row=0, column=0, sticky="w", padx=10, pady=8)
        entry_destinatario = ctk.CTkEntry(frame_campos, width=380)
        entry_destinatario.grid(row=0, column=1, padx=10, pady=8)
        ctk.CTkLabel(frame_campos, text="Cuerpo del Mensaje:").grid(row=2, column=0, sticky="w", padx=10, pady=8)
        entry_cuerpo = ctk.CTkEntry(frame_campos, width=380)
        entry_cuerpo.grid(row=2, column=1, padx=10, pady=8)

        ctk.CTkLabel(frame_campos, text="Asunto:").grid(row=1, column=0, sticky="w", padx=10, pady=8)
        entry_asunto = ctk.CTkEntry(frame_campos, width=380)
        entry_asunto.insert(0, "Oficios Generados")
        entry_asunto.grid(row=1, column=1, padx=10, pady=8)

        # Barra de progreso
        progress_label = ctk.CTkLabel(frame_campos, text="", text_color="blue")
        progress_label.grid(row=2, column=0, columnspan=2, pady=10)
        
        def enviar():
            try:
                destinatario = entry_destinatario.get().strip()
                asunto = entry_asunto.get().strip()
                cuerpo = entry_cuerpo.get().strip()
                
                if not destinatario:
                    messagebox.showwarning("Aviso", "Debes ingresar un destinatario")
                    return

                if not cuerpo:
                    messagebox.showwarning("Aviso", "Debes ingresar un cuerpo de mensaje")
                    return

                progress_label.configure(text="⏳ Generando archivos...")
                ventana.update()
                
                # Generar oficios
                formato = self.formato_var.get()
                if formato == "TXT":
                    archivos = generar_txt(personas_seleccionadas)
                elif formato == "DOCX":
                    archivos = generar_docx(personas_seleccionadas)
                elif formato == "PDF":
                    archivos = generar_pdf(personas_seleccionadas)

                progress_label.configure(text="Autenticando con Google...", text_color="white")
                ventana.update()
                
                # Enviar email usando OAuth2
                exito, mensaje = enviar_email_oauth2(destinatario, asunto, cuerpo, archivos)
                
                if exito:
                    ventana.destroy()
                    messagebox.showinfo("✅ Éxito", 
                        f"Email enviado correctamente\n\n"
                        f"Destinatario: {destinatario}\n"
                        f"Archivos adjuntos: {len(archivos)}\n\n"
                        f"{mensaje}")
                else:
                    progress_label.configure(text="❌ Error", text_color="red")
                    messagebox.showerror("Error", mensaje)
                    
            except Exception as e:
                progress_label.configure(text="❌ Error", text_color="red")
                messagebox.showerror("Error", f"Error inesperado: {str(e)}")
        
        # Botón enviar
        ctk.CTkButton(frame_campos, text="📧 Enviar por Gmail", command=enviar, 
                     height=40, font=("Arial", 14, "bold"),
                     fg_color="#DB4437", hover_color="#C23321").grid(
            row=3, column=0, columnspan=2, pady=15)


def main():
    """Función principal."""
    os.makedirs("res", exist_ok=True)
    os.makedirs(CARPETA_SALIDA, exist_ok=True)
    
    if not os.path.exists(ARCHIVO_INFO):
        with open(ARCHIVO_INFO, "w", encoding="utf-8", newline='') as f:
            escritor = csv.writer(f)
            escritor.writerow(['registro', 'nombre', 'fechaNacimiento', 'calle', 'numeroExt', 
                             'numeroInt', 'colonia', 'codigoPostal', 'ciudad', 'estado'])
    
    if not os.path.exists(ARCHIVO_PLANTILLA):
        plantilla_ejemplo = """A QUIEN CORRESPONDA:

Por medio de la presente se hace constar que {nombre}, con fecha de nacimiento {fechaNacimiento},
de {edad} años de edad, reside en:

Calle: {calle} No. Ext. {numeroExt}{numeroIntTexto}
Colonia: {colonia}
C.P. {codigoPostal}
{ciudad}, {estado}

Se extiende la presente en {ciudad}, {estado}, a {fechaActual}.

ATENTAMENTE
_____________________________________
[Firma]
"""
        with open(ARCHIVO_PLANTILLA, "w", encoding="utf-8") as f:
            f.write(plantilla_ejemplo)
    
    app = AplicacionOficios()
    app.mainloop()


if __name__ == "__main__":
    main()